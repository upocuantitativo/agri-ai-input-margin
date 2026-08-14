# -*- coding: utf-8 -*-
"""
15_build_docx_eint.py

Builds the Economics of Innovation and New Technology (Taylor & Francis)
submission DOCX from article/EINT_submission/manuscript_EINT.md.

House rules enforced here:
  * every run is BLACK (including headings, which Word styles colour blue by default)
  * no bold anywhere (table header rows get grey shading instead)
  * Times New Roman 12 pt, double spaced, numbered lines off
  * figures embedded after the paragraph of their first mention, caption below
  * tables rendered as real Word tables with a caption paragraph above

Usage:  python analysis/15_build_docx_eint.py
"""
import os, re, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
SUB = os.path.join(HERE, "..", "article", "EINT_submission")
MD = os.path.join(SUB, "manuscript_EINT.md")
FIGDIR = os.path.join(SUB, "figures")
OUT = os.path.join(SUB, "01_Manuscript_EINT.docx")

BLACK = RGBColor(0x00, 0x00, 0x00)
FONT = "Times New Roman"
SIZE = Pt(12)


def shade(cell, hexcolour="D9D9D9"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolour)
    tcPr.append(shd)


def blacken(run, size=SIZE, italic=None):
    run.font.color.rgb = BLACK
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = size
    run.bold = False
    if italic is not None:
        run.italic = italic


def add_runs(p, text, size=SIZE):
    """Italics only (*...*); ** ** is stripped to plain text, never bold."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    for part in re.split(r"(\*[^*]+?\*)", text):
        if not part:
            continue
        if part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = p.add_run(part[1:-1]); blacken(r, size, italic=True)
        else:
            r = p.add_run(part); blacken(r, size, italic=False)
    return p


def new_para(doc, text="", size=SIZE, space_after=Pt(0), spacing=2.0,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE if spacing == 2.0 else WD_LINE_SPACING.SINGLE
    if spacing != 2.0:
        pf.line_spacing = spacing
    pf.space_after = space_after
    p.alignment = align
    if indent is not None:
        pf.left_indent = indent
    if text:
        add_runs(p, text, size)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    size = {0: Pt(16), 1: Pt(14), 2: Pt(13), 3: Pt(12)}[level]
    r = p.add_run(re.sub(r"\*\*(.+?)\*\*", r"\1", text))
    blacken(r, size, italic=(level == 3))
    p.style = doc.styles["Normal"]
    return p


def parse_table(doc, block):
    rows = [ln for ln in block if ln.strip().startswith("|")]
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    cells = [c for c in cells
             if not all(re.fullmatch(r":?-{2,}:?", (x or "-").strip()) for x in c)]
    if not cells:
        return
    ncol = max(len(r) for r in cells)
    t = doc.add_table(rows=len(cells), cols=ncol)
    t.style = "Table Grid"
    t.autofit = True
    for i, row in enumerate(cells):
        for j in range(ncol):
            cell = t.rows[i].cells[j]
            cell.text = ""
            pr = cell.paragraphs[0]
            pr.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pr.paragraph_format.space_after = Pt(0)
            add_runs(pr, row[j] if j < len(row) else "", size=Pt(10))
            if i == 0:
                shade(cell)
    new_para(doc, "", space_after=Pt(6))


def main():
    with open(MD, encoding="utf-8") as f:
        raw = f.read()

    # pull the figure captions out so they can be placed under each image
    captions = {}
    mcap = re.search(r"## Figure captions\n(.*)$", raw, flags=re.S)
    if mcap:
        for m in re.finditer(r"^(Figure (\d+)\..*)$", mcap.group(1), flags=re.M):
            captions[int(m.group(2))] = m.group(1).strip()

    lines = raw.split("\n")
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = SIZE
    st.font.color.rgb = BLACK
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(1)
        s.top_margin = s.bottom_margin = Inches(1)

    placed = set()
    table_block, in_table, in_captions = [], False, False
    i = 0
    while i < len(lines):
        ln = lines[i]
        s = ln.strip()

        if s.startswith("|"):
            table_block.append(ln); in_table = True; i += 1; continue
        if in_table:
            parse_table(doc, table_block); table_block, in_table = [], False

        if s.startswith("## Figure captions"):
            in_captions = True

        if s.startswith("### "):
            add_heading(doc, s[4:], 3)
        elif s.startswith("## "):
            add_heading(doc, s[3:], 2)
        elif s.startswith("# "):
            add_heading(doc, s[2:], 0)
        elif s.startswith("- "):
            p = new_para(doc, s[2:], indent=Inches(0.3))
        elif s in ("---", ""):
            pass
        elif s.startswith("Table ") and re.match(r"^Table \d+\.", s):
            new_para(doc, s, size=Pt(11), spacing=1.0, space_after=Pt(3),
                     align=WD_ALIGN_PARAGRAPH.LEFT)
        elif s.startswith("Note:"):
            p = new_para(doc, s, size=Pt(10), spacing=1.0, space_after=Pt(10),
                         align=WD_ALIGN_PARAGRAPH.LEFT)
            for r in p.runs:
                r.italic = True
        else:
            new_para(doc, s)
            # embed each figure right after the paragraph that first mentions it
            if not in_captions:
                for n in sorted(captions):
                    if n in placed:
                        continue
                    if re.search(rf"Figure {n}\b", s):
                        path = os.path.join(FIGDIR, f"Figure{n}.png")
                        if os.path.exists(path):
                            fp = doc.add_paragraph()
                            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            fp.paragraph_format.space_before = Pt(10)
                            fp.add_run().add_picture(path, width=Inches(6.0))
                            cp = new_para(doc, captions[n], size=Pt(10), spacing=1.0,
                                          space_after=Pt(12),
                                          align=WD_ALIGN_PARAGRAPH.LEFT)
                            placed.add(n)
                        break
        i += 1

    if in_table and table_block:
        parse_table(doc, table_block)

    doc.save(OUT)

    # ---- report -------------------------------------------------------
    nblue = nbold = 0
    for p in doc.paragraphs:
        for r in p.runs:
            try:
                if r.font.color is not None and r.font.color.rgb is not None \
                        and str(r.font.color.rgb) != "000000":
                    nblue += 1
            except Exception:
                pass
            if r.bold:
                nbold += 1
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    for r in p.runs:
                        if r.bold:
                            nbold += 1
    print("DOCX saved:", os.path.abspath(OUT))
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} "
          f"figures_embedded={sorted(placed)}")
    print(f"non-black runs={nblue}  bold runs={nbold}  (both must be 0)")
    missing = sorted(set(captions) - placed)
    if missing:
        print("WARNING: figures with a caption but no anchor in the text:", missing)


if __name__ == "__main__":
    main()
