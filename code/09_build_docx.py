# -*- coding: utf-8 -*-
# ---------------------------------------------------------------------------
# Replication package: agri-ai-input-margin
# Pipeline step aux -- Manuscript assembly (Spanish v1). Requires article/manuscrito.md, which is NOT shipped here.
#
# Copied verbatim from the author's working tree (analysis/09_build_docx.py); this header
# comment is the only modification. Run from the REPOSITORY ROOT, e.g.
#     python code/09_build_docx.py
# because some scripts resolve paths relative to the working directory and
# others relative to their own location; both conventions resolve correctly
# only when the working directory is the repository root.
# ---------------------------------------------------------------------------
"""
09_build_docx.py
Convierte article/manuscrito.md a un DOCX revisable (encabezados, tablas, citas,
negritas y un apéndice de figuras embebidas). Pragmatico pero fiel al contenido.
"""
import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
ART = os.path.join(HERE, "..", "article")
FIG = os.path.join(HERE, "..", "figures")
md_path = os.path.join(ART, "manuscrito.md")
out_path = os.path.join(ART, "manuscrito.docx")

with open(md_path, encoding="utf-8") as f:
    lines = f.read().split("\n")

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"; style.font.size = Pt(11)

# mapa figura -> caption
fig_caps = {
    1: ("fig1_produccion_mundial.png", "Producción agroalimentaria mundial por gran grupo de cultivo (FAOSTAT)."),
    2: ("fig2_desaceleracion_rendimiento.png", "Desaceleración del crecimiento del rendimiento cerealista mundial."),
    3: ("fig3_mecanizacion_espana.png", "España: hectáreas de tierra arable por tractor (1961-2009)."),
    4: ("fig4_nitrogeno.png", "Rendimiento vs nitrógeno y productividad parcial decreciente del N."),
    5: ("fig5_convergencia.png", "σ-convergencia y β-convergencia de los rendimientos entre países."),
    6: ("fig6_rendimiento_por_renta.png", "Rendimiento cerealista mediano por grupo de renta."),
    7: ("fig7_importancia.png", "Importancia por permutación del modelo de gradient boosting."),
    8: ("fig8_generico_vs_especifico.png", "Error de predicción: genérico vs específico."),
    9: ("fig9_empleo.png", "Empleo agrícola mundial y relación con la productividad del trabajo."),
    10: ("fig10_dependencia_parcial.png", "Dependencia parcial: rendimientos marginales decrecientes de los insumos."),
}

def add_runs(p, text):
    # negritas **...**
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        else:
            r = p.add_run(part)

def parse_table(block):
    rows = [ln for ln in block if ln.strip().startswith("|")]
    # quitar separador |---|
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    cells = [c for c in cells if not all(re.fullmatch(r":?-{2,}:?", x or "-") for x in c)]
    if not cells:
        return
    t = doc.add_table(rows=len(cells), cols=len(cells[0]))
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(cells):
        for j, val in enumerate(row):
            if j < len(t.rows[i].cells):
                cell = t.rows[i].cells[j]
                cell.text = ""
                pr = cell.paragraphs[0]
                add_runs(pr, val.replace("**", ""))
                if i == 0:
                    for rr in pr.runs:
                        rr.bold = True
    doc.add_paragraph("")

i = 0
table_block = []
in_table = False
while i < len(lines):
    ln = lines[i]
    stripped = ln.strip()
    # tablas
    if stripped.startswith("|"):
        table_block.append(ln); in_table = True; i += 1; continue
    elif in_table:
        parse_table(table_block); table_block = []; in_table = False
    # encabezados
    if stripped.startswith("### "):
        doc.add_heading(stripped[4:], level=3)
    elif stripped.startswith("## "):
        doc.add_heading(stripped[3:], level=2)
    elif stripped.startswith("# "):
        h = doc.add_heading(stripped[2:], level=0)
    elif stripped.startswith("> "):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.4)
        r = p.add_run(stripped[2:].replace("**", "")); r.italic = True
    elif stripped == "---" or stripped == "":
        if stripped == "":
            pass
    else:
        p = doc.add_paragraph()
        add_runs(p, stripped)
    i += 1
if in_table and table_block:
    parse_table(table_block)

# apéndice de figuras
doc.add_page_break()
doc.add_heading("Apéndice de figuras", level=1)
for n in sorted(fig_caps):
    fn, cap = fig_caps[n]
    fp = os.path.join(FIG, fn)
    if os.path.exists(fp):
        doc.add_picture(fp, width=Inches(6.0))
        last = doc.paragraphs[-1]; last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp = doc.add_paragraph(); r = cp.add_run(f"Figura {n}. {cap}")
        r.italic = True; r.font.size = Pt(9)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(out_path)
print("DOCX guardado:", out_path)
print("Parrafos:", len(doc.paragraphs), "| Tablas:", len(doc.tables))
