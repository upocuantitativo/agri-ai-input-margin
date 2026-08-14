# -*- coding: utf-8 -*-
# ---------------------------------------------------------------------------
# Replication package: agri-ai-input-margin
# Pipeline step aux -- Manuscript assembly (English v2). Requires article/manuscrito_EN_v2.md, which is NOT shipped here.
#
# Copied verbatim from the author's working tree (analysis/13_build_docx_v2.py); this header
# comment is the only modification. Run from the REPOSITORY ROOT, e.g.
#     python code/13_build_docx_v2.py
# because some scripts resolve paths relative to the working directory and
# others relative to their own location; both conventions resolve correctly
# only when the working directory is the repository root.
# ---------------------------------------------------------------------------
"""
13_build_docx_v2.py

Convierte article/manuscrito_EN_v2.md en un DOCX revisable en el que TODO EL
TEXTO NUEVO O REESCRITO APARECE EN AZUL, para que el editor y los revisores de
la nueva revista vean de un vistazo que se ha respondido a la critica.

Convencion en el markdown: el texto entre {{ y }} se renderiza en azul.
Soporta encabezados, parrafos, citas en bloque, listas, negritas, cursivas y
tablas (incluida la coloracion de celdas completas o de fragmentos).
"""
import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
ART = os.path.join(HERE, "..", "article")
md_path = os.path.join(ART, "manuscrito_EN_v2.md")
out_path = os.path.join(ART, "manuscrito_EN_v2.docx")

BLUE = RGBColor(0x00, 0x00, 0xCC)

with open(md_path, encoding="utf-8") as f:
    lines = f.read().split("\n")

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def add_runs(p, text, force_blue=False):
    """Anade runs aplicando negrita/cursiva y azul a los tramos entre {{ }}."""
    for chunk, blue in split_blue(text):
        blue = blue or force_blue
        for part in re.split(r"(\*\*.+?\*\*|\*[^*]+?\*)", chunk):
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                r = p.add_run(part[2:-2]); r.bold = True
            elif part.startswith("*") and part.endswith("*") and len(part) > 2:
                r = p.add_run(part[1:-1]); r.italic = True
            else:
                r = p.add_run(part)
            if blue:
                r.font.color.rgb = BLUE


def split_blue(text):
    """Devuelve [(fragmento, es_azul), ...] segun los marcadores {{ }}."""
    out, pos = [], 0
    for m in re.finditer(r"\{\{(.*?)\}\}", text, flags=re.S):
        if m.start() > pos:
            out.append((text[pos:m.start()], False))
        out.append((m.group(1), True))
        pos = m.end()
    if pos < len(text):
        out.append((text[pos:], False))
    return out or [(text, False)]


def parse_table(block):
    rows = [ln for ln in block if ln.strip().startswith("|")]
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    cells = [c for c in cells
             if not all(re.fullmatch(r":?-{2,}:?", x or "-") for x in c)]
    if not cells:
        return
    t = doc.add_table(rows=len(cells), cols=len(cells[0]))
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(cells):
        for j, val in enumerate(row):
            if j >= len(t.rows[i].cells):
                continue
            cell = t.rows[i].cells[j]
            cell.text = ""
            pr = cell.paragraphs[0]
            add_runs(pr, val)
            if i == 0:
                for rr in pr.runs:
                    rr.bold = True
    doc.add_paragraph("")


i, table_block, in_table = 0, [], False
while i < len(lines):
    ln = lines[i]
    s = ln.strip()
    if s.startswith("|"):
        table_block.append(ln); in_table = True; i += 1; continue
    elif in_table:
        parse_table(table_block); table_block, in_table = [], False

    if s.startswith("### "):
        h = doc.add_heading("", level=3); add_runs(h, s[4:])
    elif s.startswith("## "):
        h = doc.add_heading("", level=2); add_runs(h, s[3:])
    elif s.startswith("# "):
        h = doc.add_heading("", level=0); add_runs(h, s[2:])
    elif s.startswith("> "):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.4)
        add_runs(p, s[2:])
        for r in p.runs:
            r.italic = True
    elif re.match(r"^\d+\.\s", s) or s.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        add_runs(p, re.sub(r"^(\d+\.\s|-\s)", "", s))
    elif s in ("---", ""):
        pass
    else:
        p = doc.add_paragraph(); add_runs(p, s)
    i += 1

if in_table and table_block:
    parse_table(table_block)

doc.save(out_path)

# recuento de palabras aproximado y de tramos en azul
txt = "\n".join(lines)
blue_words = sum(len(m.split()) for m in re.findall(r"\{\{(.*?)\}\}", txt, flags=re.S))
total_words = len(re.sub(r"[{}]", "", txt).split())
print("DOCX guardado:", out_path)
print(f"Parrafos: {len(doc.paragraphs)} | Tablas: {len(doc.tables)}")
print(f"Palabras totales ~{total_words} | en azul ~{blue_words} "
      f"({100*blue_words/max(total_words,1):.0f}% del texto es nuevo)")
