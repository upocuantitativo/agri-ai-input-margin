# -*- coding: utf-8 -*-
"""
16_build_docx_aux.py

Renders the auxiliary EINT submission documents (title page, cover letter,
supplementary material) from markdown to plain black, non-bold DOCX.

Usage:  python analysis/16_build_docx_aux.py
"""
import os, re, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
SUB = os.path.join(HERE, "..", "article", "EINT_submission")
BLACK = RGBColor(0, 0, 0)
FONT = "Times New Roman"

JOBS = [
    ("title_page.md", "02_Title_page_EINT.docx"),
    ("cover_letter.md", "03_Cover_letter_EINT.docx"),
    ("supplementary.md", "04_Supplementary_material_EINT.docx"),
]


def shade(cell, hexcolour="D9D9D9"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolour)
    tcPr.append(shd)


def blacken(run, size, italic=False):
    run.font.color.rgb = BLACK
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = size
    run.bold = False
    run.italic = italic


def add_runs(p, text, size=Pt(12)):
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    for part in re.split(r"(\*[^*]+?\*)", text):
        if not part:
            continue
        if part.startswith("*") and part.endswith("*") and len(part) > 2:
            blacken(p.add_run(part[1:-1]), size, italic=True)
        else:
            blacken(p.add_run(part), size)


def parse_table(doc, block):
    rows = [ln for ln in block if ln.strip().startswith("|")]
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    cells = [c for c in cells
             if not all(re.fullmatch(r":?-{2,}:?", (x or "-").strip()) for x in c)]
    if not cells:
        return
    ncol = max(len(r) for r in cells)
    t = doc.add_table(rows=len(cells), cols=ncol)
    t.style = "Table Grid"
    for i, row in enumerate(cells):
        for j in range(ncol):
            cell = t.rows[i].cells[j]
            cell.text = ""
            pr = cell.paragraphs[0]
            pr.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pr.paragraph_format.space_after = Pt(0)
            add_runs(pr, row[j] if j < len(row) else "", size=Pt(9))
            if i == 0:
                shade(cell)
    doc.add_paragraph("")


def build(src, dst, figdir=None):
    path = os.path.join(SUB, src)
    if not os.path.exists(path):
        print("skip (not found):", src)
        return
    lines = open(path, encoding="utf-8").read().split("\n")
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(12)
    st.font.color.rgb = BLACK
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(1)

    table_block, in_table = [], False
    for ln in lines:
        s = ln.strip()
        if s.startswith("|"):
            table_block.append(ln); in_table = True; continue
        if in_table:
            parse_table(doc, table_block); table_block, in_table = [], False

        if s.startswith("!["):
            m = re.match(r"!\[[^\]]*\]\(([^)]+)\)", s)
            if m and figdir:
                img = os.path.join(figdir, os.path.basename(m.group(1)))
                if os.path.exists(img):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(img, width=Inches(6.0))
            continue

        lvl = len(s) - len(s.lstrip("#"))
        if lvl and s[lvl:lvl + 1] == " ":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            blacken(p.add_run(re.sub(r"\*\*(.+?)\*\*", r"\1", s[lvl + 1:])),
                    {1: Pt(16), 2: Pt(13), 3: Pt(12)}.get(lvl, Pt(12)),
                    italic=(lvl >= 3))
        elif s in ("---", ""):
            doc.add_paragraph("")
        elif s.startswith("- "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_runs(p, "• " + s[2:])
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            size = Pt(9) if s.startswith("Note:") else Pt(12)
            add_runs(p, s, size)

    if in_table and table_block:
        parse_table(doc, table_block)

    out = os.path.join(SUB, dst)
    doc.save(out)
    bad = sum(1 for p in doc.paragraphs for r in p.runs if r.bold)
    print(f"saved {dst}  paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} bold_runs={bad}")


if __name__ == "__main__":
    figdir = os.path.join(SUB, "figures", "supplementary")
    for src, dst in JOBS:
        build(src, dst, figdir=figdir)
