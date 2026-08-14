# -*- coding: utf-8 -*-
# ---------------------------------------------------------------------------
# Replication package: agri-ai-input-margin
# Pipeline step aux -- Manuscript assembly (English v1). Requires article/manuscrito_EN.md, which is NOT shipped here.
#
# Copied verbatim from the author's working tree (analysis/09b_build_docx_en.py); this header
# comment is the only modification. Run from the REPOSITORY ROOT, e.g.
#     python code/09b_build_docx_en.py
# because some scripts resolve paths relative to the working directory and
# others relative to their own location; both conventions resolve correctly
# only when the working directory is the repository root.
# ---------------------------------------------------------------------------
"""
09b_build_docx_en.py
Convierte article/manuscrito_EN.md a un DOCX revisable, SIN apéndice de figuras
(el usuario pidió la traducción sin gráficos). Encabezados, tablas, negritas y citas.
"""
import os, re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
ART = os.path.join(HERE, "..", "article")
md_path = os.path.join(ART, "manuscrito_EN.md")
out_path = os.path.join(ART, "manuscrito_EN.docx")

with open(md_path, encoding="utf-8") as f:
    lines = f.read().split("\n")

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"; style.font.size = Pt(11)

def add_runs(p, text):
    parts = re.split(r"(\*\*.+?\*\*|\*.+?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = p.add_run(part[1:-1]); r.italic = True
        else:
            p.add_run(part)

def parse_table(block):
    rows = [ln for ln in block if ln.strip().startswith("|")]
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    cells = [c for c in cells if not all(re.fullmatch(r":?-{2,}:?", x or "-") for x in c)]
    if not cells:
        return
    t = doc.add_table(rows=len(cells), cols=len(cells[0]))
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(cells):
        for j, val in enumerate(row):
            if j < len(t.rows[i].cells):
                cell = t.rows[i].cells[j]
                cell.text = ""
                pr = cell.paragraphs[0]
                add_runs(pr, val.replace("**", ""))
                if i == 0:
                    for rr in pr.runs:
                        rr.bold = True
    doc.add_paragraph("")

i = 0
table_block = []
in_table = False
while i < len(lines):
    ln = lines[i]
    stripped = ln.strip()
    if stripped.startswith("|"):
        table_block.append(ln); in_table = True; i += 1; continue
    elif in_table:
        parse_table(table_block); table_block = []; in_table = False
    if stripped.startswith("### "):
        doc.add_heading(stripped[4:], level=3)
    elif stripped.startswith("## "):
        doc.add_heading(stripped[3:], level=2)
    elif stripped.startswith("# "):
        doc.add_heading(stripped[2:], level=0)
    elif stripped.startswith("> "):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.4)
        add_runs(p, stripped[2:]);
        for r in p.runs: r.italic = True
    elif stripped == "---" or stripped == "":
        pass
    else:
        p = doc.add_paragraph()
        add_runs(p, stripped)
    i += 1
if in_table and table_block:
    parse_table(table_block)

doc.save(out_path)
print("DOCX guardado:", out_path)
print("Parrafos:", len(doc.paragraphs), "| Tablas:", len(doc.tables))
